from __future__ import annotations

import csv
import mimetypes
import tempfile
from pathlib import Path
from typing import BinaryIO

from open_notebook.models import SourceKind, SourceOut
from open_notebook.services.storage import Storage


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".json", ".csv"}
DOC_SUFFIXES = {".pdf", ".docx"}
SHEET_SUFFIXES = {".xlsx", ".xlsm", ".xls", ".csv"}


def infer_source_kind(filename: str) -> SourceKind:
    suffix = Path(filename).suffix.lower()
    if suffix in IMAGE_SUFFIXES:
        return SourceKind.image
    if suffix in SHEET_SUFFIXES:
        return SourceKind.spreadsheet
    if suffix in DOC_SUFFIXES or suffix in {".md", ".txt"}:
        return SourceKind.document
    return SourceKind.unknown


async def save_upload(
    storage: Storage,
    *,
    session_id: str,
    filename: str,
    fileobj: BinaryIO,
) -> SourceOut:
    suffix = Path(filename).suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        while chunk := fileobj.read(1024 * 1024):
            tmp.write(chunk)
        tmp_path = Path(tmp.name)
    try:
        kind = infer_source_kind(filename)
        summary, metadata = parse_source(tmp_path, filename, kind)
        return storage.save_source(
            session_id=session_id,
            filename=filename,
            raw_path=tmp_path,
            kind=kind,
            summary=summary,
            metadata=metadata,
        )
    finally:
        tmp_path.unlink(missing_ok=True)


def parse_source(path: Path, filename: str, kind: SourceKind) -> tuple[str, dict]:
    mime = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    metadata = {"mime_type": mime, "bytes": path.stat().st_size}
    suffix = path.suffix.lower()
    try:
        if kind == SourceKind.image:
            from PIL import Image

            with Image.open(path) as img:
                metadata.update({"width": img.width, "height": img.height, "mode": img.mode})
            return f"Image {filename}: {metadata.get('width')}x{metadata.get('height')}.", metadata
        if suffix in {".txt", ".md", ".markdown", ".json"}:
            text = path.read_text(encoding="utf-8", errors="replace")
            metadata["text_excerpt"] = text[:12000]
            return _summarize_text(text, filename), metadata
        if suffix == ".csv":
            return _parse_csv(path, filename, metadata)
        if suffix in {".xlsx", ".xlsm"}:
            return _parse_xlsx(path, filename, metadata)
        if suffix == ".docx":
            return _parse_docx(path, filename, metadata)
        if suffix == ".pdf":
            return _parse_pdf(path, filename, metadata)
    except Exception as exc:
        metadata["parse_error"] = str(exc)
    return f"{filename}: uploaded file, {metadata['bytes']} bytes.", metadata


def _summarize_text(text: str, filename: str) -> str:
    compact = " ".join(text.split())
    if len(compact) > 900:
        compact = compact[:900] + "..."
    return f"{filename}: {compact}" if compact else f"{filename}: empty text file."


def _parse_csv(path: Path, filename: str, metadata: dict) -> tuple[str, dict]:
    rows = []
    with path.open(newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for idx, row in enumerate(reader):
            rows.append(row)
            if idx >= 20:
                break
    metadata["preview_rows"] = rows
    metadata["text_excerpt"] = "\n".join(",".join(r) for r in rows)
    header = rows[0] if rows else []
    return f"{filename}: CSV table with columns {', '.join(header[:12])}.", metadata


def _parse_xlsx(path: Path, filename: str, metadata: dict) -> tuple[str, dict]:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets = []
    previews = {}
    for ws in wb.worksheets[:8]:
        rows = []
        for idx, row in enumerate(ws.iter_rows(values_only=True)):
            rows.append([str(x) if x is not None else "" for x in row[:20]])
            if idx >= 15:
                break
        sheets.append({"name": ws.title, "max_row": ws.max_row, "max_column": ws.max_column})
        previews[ws.title] = rows
    metadata["sheets"] = sheets
    metadata["preview_rows"] = previews
    metadata["text_excerpt"] = str(previews)[:12000]
    return f"{filename}: spreadsheet with sheets " + ", ".join(s["name"] for s in sheets) + ".", metadata


def _parse_docx(path: Path, filename: str, metadata: dict) -> tuple[str, dict]:
    from docx import Document

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    metadata["paragraphs"] = len(doc.paragraphs)
    metadata["text_excerpt"] = text[:12000]
    return _summarize_text(text, filename), metadata


def _parse_pdf(path: Path, filename: str, metadata: dict) -> tuple[str, dict]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages[:30]:
        parts.append(page.extract_text() or "")
    text = "\n".join(parts)
    metadata["pages"] = len(reader.pages)
    metadata["text_excerpt"] = text[:12000]
    return _summarize_text(text, filename), metadata
